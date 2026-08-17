import copy
import os
import docx
from docx.shared import Pt

SRC = r"C:\Users\Probook\NCIRL 2025-2026\NCIRL CEAI\CA 3\x25128442_Sydorenko_Nadiya_CEAI_CA3.docx"
DST = r"C:\Users\Probook\NCIRL 2025-2026\NCIRL CEAI\CA 3\x25128442_Sydorenko_Nadiya_CEAI_CA3.docx"


from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_para_text(p, text):
    for child in list(p._element):
        if child.tag != qn("w:pPr"):
            p._element.remove(child)
    p.add_run(text)


doc = docx.Document(SRC)

paras = list(doc.paragraphs)
delete_from = None
for i, p in enumerate(paras):
    if p.style.name == "Heading 1" and p.text.strip().startswith("Introduction"):
        delete_from = i
        break

for p in paras[delete_from:]:
    p._element.getparent().remove(p._element)

for p in doc.paragraphs:
    t = p.text.strip()
    if "Individual Project: Build It Live, Prove It's Real" in t:
        set_para_text(p, "Individual Project: Build an Agentic Organisation")
    elif t.startswith("Scenario Code"):
        set_para_text(p, "Organisation: Prism — an Agentic Crypto Research Copilot")
    elif "ironforge-chatbot" in t or "IronForge" in t:
        set_para_text(p, "Live GitHub Pages URL: https://nadiya-nci.github.io/prism/")

def add(text, style="Normal", bold=None):
    p = doc.add_paragraph()
    p.style = doc.styles[style]
    r = p.add_run(text)
    if bold is not None:
        r.bold = bold
    return p

def add_heading(text):
    add(text, "Heading 1")

def add_subheading(text):
    add(text, "Heading 2")

def _boxed(p):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "8")
        e.set(qn("w:space"), "4")
        e.set(qn("w:color"), "6C7CFF")
        pBdr.append(e)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "EDEFFF")
    pPr.append(shd)

def add_fig(n, title, what):
    p = doc.add_paragraph()
    _boxed(p)
    r = p.add_run("[ INSERT FIGURE " + str(n) + " HERE ]")
    r.bold = True
    p.add_run("\n" + title)
    cap = doc.add_paragraph()
    cr = cap.add_run("Figure " + str(n) + ". " + title + ". " + what)
    cr.italic = True

add_heading("Introduction")
add(
    "This report documents the design, build, and live demonstration of Prism, a fully agentic "
    "organisation of five specialised AI agents that serves retail crypto investors with trustworthy, "
    "live-sourced research. Each agent embodies one innovation archetype — Researcher, Designer, Maker, "
    "Communicator and Manager — and hands its complete output to the next agent in a fixed pipeline. "
    "The organisation produced, end to end, an opportunity brief, a product concept, a working prototype "
    "deployed on GitHub Pages, a go-to-market plan, and an executive summary. At least one agent connects to "
    "a live external data source at query time: the Researcher, Designer, Maker, and Manager all execute live "
    "tool calls against CoinGecko, a user-owned Google Sheet, Alternative.me and the open ExchangeRate-API. "
    "No API key or credential is committed anywhere in the repository. Section: Your Organisation follows."
)

add_heading("1. Your Organisation")
add(
    "Prism is a fictional, AI-native venture organisation serving real retail-investor archetypes — people "
    "who bought digital assets near market peaks, are frequently underwater, and distrust the hype-driven "
    "content around them. Its business challenge is customer engagement in a fear-driven market: attracting "
    "anxious investors, engaging them with evidence, persuading them with grounded answers, and retaining them "
    "through honesty. Live data observed during this project run — a watchlist portfolio valued at roughly "
    "−61% versus book cost, with the Fear & Greed index at 34 (Fear) — is exactly the anxiety this "
    "organisation is built to address."
)
add(
    "The challenge benefits from an agentic approach because no single agent could research live markets, "
    "design a product, ship working code, write persuasively and govern strategically at the same time. The "
    "pipeline Researcher \u2192 Designer \u2192 Maker \u2192 Communicator \u2192 Manager makes work cumulative: each "
    "stage ingests the previous stage's full deliverable, so the final output genuinely exceeds what one "
    "agent could produce alone. The live-data requirement reinforces the customer-engagement argument: a "
    "research copilot is only as trustworthy as the freshness it can display in real time."
)

add_heading("2. Agent Designs")
add(
    "Members intervene in order; each has a distinct system prompt, personality and domain expertise, and each "
    "is wired to specific live data tools. Full system prompts are reproduced below from pipeline/agent_prompts.py."
)

agents = [
    (
        "2.1 Vera Vector — Researcher (archetype: deep analyst)",
        "Personality: rigorous, numbers-first, mildly sceptical of hype. Reads live data before speaking and "
        "refuses invented statistics. Tool calls: get_live_market (CoinGecko), get_market_sentiment "
        "(Alternative.me), get_exchange_rate (live FX, ExchangeRate-API). Produces: a research brief with live "
        "snapshot, movers, three opportunity hypotheses and risk flags.",
        "System prompt: \"You are Vera Vector, HEAD OF MARKET RESEARCH at Prism, an agentic organisation "
        "serving retail crypto investors. You are rigorous, numbers-first and mildly sceptical of hype. Your "
        "superpower is deep analysis and pattern recognition; you read live data before you speak. You ALWAYS "
        "call tools to pull live market data and sentiment, and you DARE NOT invent statistics. Style: precise, "
        "structured, cautious optimism. Your deliverable is a RESEARCH BRIEF in markdown with these sections: "
        "1) Macro & sentiment snapshot (from live tool calls, with figures); 2) Movers worth attention; 3) 3 "
        "concrete opportunity hypotheses testable by the rest of the organisation; 4) Risk flags every later "
        "agent must respect.\"",
    ),
    (
        "2.2 Maya Möbius — Designer (archetype: creative strategist)",
        "Personality: imaginative, user-obsessed, allergic to clutter. Tool call: get_portfolio_sheet (live "
        "Google Sheet watchlist) so the concept fits a real portfolio. Produces: a solution concept with persona, "
        "journey, feature spec, success metrics and constraints.",
        "System prompt: \"You are Maya Möbius, HEAD OF PRODUCT DESIGN at Prism. You are imaginative, "
        "user-obsessed, and allergic to clutter. Your superpower is creative problem-solving and design "
        "thinking: you turn raw research into an experience real people want. You ALWAYS call your tool to read "
        "the user's live portfolio watchlist so your concept fits their reality. Style: bold ideas, plain "
        "language, actionable specs. Given the Researcher's brief, produce a SOLUTION CONCEPT in markdown: "
        "1) Target persona described concretely; 2) Core user journey; 3) Feature specification (exactly what "
        "the prototype must include); 4) Success metrics; 5) Constraints the Maker must respect.\"",
    ),
    (
        "2.3 Code Flint — Maker (archetype: craftsman)",
        "Personality: pragmatic, fast, allergic to scope creep; claims only what the build can prove. Tool "
        "calls: get_portfolio_sheet, get_market_sentiment to verify live sources before building. Produces: a "
        "build report mapping each feature to shipped code with live-connection proof.",
        "System prompt: \"You are Code Flint, LEAD ENGINEER at Prism. Your superpower is technical "
        "craftsmanship and rapid prototyping: you turn design into a working, verifiable artifact. You ALWAYS "
        "verify live sources with your tools so the build is grounded in real queries, never hardcoded values. "
        "Style: terse, technical, honest about trade-offs. Given the Designer's concept, produce a BUILD REPORT "
        "in markdown: 1) Tech stack & architecture (static GitHub Pages site + live fetches in browser JS); 2) "
        "What was built, mapped to each feature; 3) Live data connections implemented (queried at runtime, no "
        "keys committed); 4) How the pipeline's agents hand work to each other; 5) Test evidence and known "
        "limitations.\"",
    ),
    (
        "2.4 Riley Rhetoric — Communicator (archetype: storyteller)",
        "Personality: magnetic storyteller who turns technology into trust; channel-aware; refuses hype. Tool "
        "call: get_market_sentiment so campaign tone matches live market mood. Produces: positioning, campaign "
        "concept, drafted channel copy and compliance guardrails.",
        "System prompt: \"You are Riley Rhetoric, CHIEF MARKETING OFFICER at Prism. Your superpower is "
        "persuasion and storytelling; you write copy that speaks to one anxious, curious investor at a time. "
        "You ALWAYS call your sentiment tool so your campaign tone matches the current market mood. Style: warm, "
        "confident, honest — never hype, never fear-mongering. Given the Maker's build report, produce a "
        "GO-TO-MARKET plan in markdown: 1) Positioning statement; 2) Campaign concept + name; 3) 3-channel "
        "rollout with actual copy drafted; 4) Conversion funnel; 5) Compliance guardrails for every asset.\"",
    ),
    (
        "2.5 Atlas — Manager (archetype: orchestrator)",
        "Personality: calm, accountable, synthesizing; asks whether the whole is coherent, valuable and "
        "trustworthy. Tool calls: get_portfolio_sheet, get_live_market, get_exchange_rate — re-fetches live "
        "data at the moment of review. Produces: executive summary with alignment review, KPIs, risk register "
        "and four-week plan.",
        "System prompt: \"You are Atlas, CEO of Prism. You are calm, accountable and synthetic: you read "
        "every teammate's output and decide whether the organisation produced something coherent, valuable and "
        "trustworthy. Your superpower is leadership and orchestration. You ALWAYS re-fetch the live portfolio "
        "sheet and market data yourself so your review reflects reality at the moment you report, and you verify "
        "the chain was unbroken. Style: executive, decisive. Given ALL previous outputs, produce an EXECUTIVE "
        "SUMMARY in markdown: 1) Strategic alignment review; 2) KPIs; 3) Risk register (GDPR, EU AI Act Article "
        "50, no-advice boundary); 4) Evidence of collaboration; 5) Next 4-week operating plan.\"",
    ),
]
for title, desc, prompt in agents:
    add(title, "Normal", bold=True)
    add(desc)
    add(prompt)

add(
    "Differentiation: the five agents vary in temperament (sceptic, idealist, pragmatist, storyteller, "
    "orchestrator), skill (analysis, design, engineering, persuasion, governance), allowed tools, and output "
    "format — so each reads and works unlike the others despite sharing one codebase."
)

add_heading("3. The Pipeline in Action")
add(
    "Each stage receives the previous stage's full output as its input (PREVIOUS STAGE OUTPUT) and adds its "
    "own layer; deliverables accumulate into one decision-ready package. Evidence, captured 16/08/2026 "
    "23:06 UTC (see output/free-run/ and transcript.json):",
)
add(
    "Researcher tool-called CoinGecko (top-10 in EUR), Alternative.me (Fear) and the ExchangeRate-API for FX, then produced the "
    "research brief. Its −61% portfolio observation became the Designer's persona (\"Denis, underwater\"). The "
    "Designer then tool-called the live Google Sheet watchlist via published CSV and produced the concept; the "
    "Maker tool-called the sheet and sentiment, then shipped the actual GitHub Pages site (live market table, "
    "watchlist valuation, mood chip, chat). The Communicator tool-called sentiment and drafted the \"See "
    "Clearly\" Fear-34 campaign. Finally the Manager re-fetched CoinGecko, the sheet and FX live, verified the "
    "same −61% figure independently, and issued the executive summary confirming the chain was unbroken.",
)
add(
    "Cumulative output: a research brief \u2192 a concept \u2192 a deployed working prototype \u2192 a go-to-market "
    "plan \u2192 an executive summary. No single agent could have produced the deployed, marketed, governed "
    "organisation alone. Each screenshot below shows the actual output of that stage as produced during the "
    "run; the live prototype is reachable at the GitHub Pages URL on the cover, and the chat worker at "
    "https://prism.n-sydorenko-mail.workers.dev."
)

add_subheading("3.1 Stage handoff evidence")
add_fig(
    1,
    "Stage 1 — Research brief (Vera Vector)",
    "output/free-run/01_researcher.md: live top-10 prices (CoinGecko), Fear & Greed index 34 (Alternative.me), "
    "FX (ExchangeRate-API), the \u221261% portfolio observation, and the three opportunity hypotheses handed to "
    "the Designer.",
)
add_fig(
    2,
    "Stage 2 — Solution concept (Maya M\u00f6bius)",
    "output/free-run/02_designer.md: the 'Denis' persona built from the Researcher's \u221261% finding, the user "
    "journey, the feature specification (F1-F5) and the constraints handed to the Maker.",
)
add_fig(
    3,
    "Stage 3 — Build report (Code Flint)",
    "output/free-run/03_maker.md: tech stack and architecture, feature-to-build mapping, live data connections "
    "(queried at runtime, no keys committed) and test evidence handed to the Communicator.",
)
add_fig(
    4,
    "Stage 4 — Go-to-market plan (Riley Rhetoric)",
    "output/free-run/04_communicator.md: the 'See Clearly' campaign, three-channel rollout with drafted copy and "
    "compliance guardrails handed to the Manager.",
)
add_fig(
    5,
    "Stage 5 — Executive summary (Atlas)",
    "output/free-run/05_manager.md: alignment review, launch KPIs, GDPR / EU AI Act risk register and the "
    "four-week operating plan.",
)
add_fig(
    6,
    "Live tool-call transcript",
    "output/free-run/transcript.json: timestamped tool_calls per stage, proving every agent queried a live "
    "external source at query time rather than using hardcoded values.",
)

add_subheading("3.2 The working prototype")
add(
    "The prototype is deployed on GitHub Pages and reachable without login at "
    "https://nadiya-nci.github.io/prism/. Data refreshes on every load and every two minutes; the market table, "
    "mood card and watchlist all carry fetch timestamps."
)
add_fig(
    7,
    "Live site, full page (top)",
    "the hero, live price ticker, market mood and portfolio cards with visible fetch timestamps, and the live "
    "market table with coin logos, ranks and sparklines.",
)
add_fig(
    8,
    "Live watchlist valuation",
    "the 'Your holdings' card computing current value against cost basis live from the published Google Sheet "
    "(approximately \u221261%).",
)
add_fig(
    9,
    "Live chat interaction",
    "one grounded question and answer in the chat copilot, showing the research-support-not-advice boundary and "
    "the live snapshot timestamp, served via the Cloudflare Worker.",
)

add_heading("4. Regulatory and Ethical Considerations")
add(
    "Under the GDPR (Regulation (EU) 2016/679), Prism collects no personal data and runs no accounts: the "
    "watchlist is the user's own sheet, read fetch-only, honouring Article 5 data minimisation. Under the EU AI "
    "Act (Regulation (EU) 2024/1689), the copilot sits adjacent to Annex III high-risk domains (financial "
    "assistance), so the organisation imposes transparency duties consistent with Article 50: AI-generated "
    "content is labelled, every figure carries its live source and timestamp, and the interface states the "
    "user is chatting with an AI, not a human. System prompts enforce a research-only boundary (no instruction "
    "verbs such as buy/sell), and the Manager's risk register tracks the advice boundary as a top control. "
    "Trust is engineered: stale data is impossible by design (runtime fetches, 2-minute refresh), credentials "
    "live only in a serverless secret store, and campaign copy avoids fear-mongering in a Fear-34 market."
)

add_heading("5. Reflection")
add(
    "[USER TO WRITE IN OWN WORDS — ~300 words, replace this paragraph. Suggested coverage: what worked, what "
    "didn't, what surprised you, what you learned about multi-agent collaboration, what you would improve with "
    "more time. Evidence: describe at least one iteration (e.g. switching the chat to a Cloudflare worker to "
    "keep the key server-side, or a second pipeline run whose numbers changed because the market moved).]"
)

add_heading("Conclusion")
add(
    "The project set out to build an agentic organisation, not merely describe one. Five specialised agents, "
    "each with its own system prompt, personality and live-data tools, worked in a fixed pipeline and produced, "
    "end to end, a research brief, a product concept, a deployed prototype, a go-to-market plan and an executive "
    "summary. Every stage queried live external data at the moment of use, no credential was committed, and the "
    "prototype remains reachable on GitHub Pages. The clearest proof that the handoff chain works is cumulative: "
    "the \u221261% portfolio finding surfaced by the Researcher survived unchanged through design, build and "
    "marketing, and was independently re-confirmed by the Manager. The pipeline produced results no single agent "
    "could have produced alone."
)

add_heading("References")
add(
    "AI usage and attribution. The five-agent pipeline was executed on a free opencode assistant model using "
    "the system prompts reproduced in Section 2 (pipeline/agent_prompts.py). The live site chat is served by "
    "Claude (claude-sonnet-4-5) through a Cloudflare Worker, keeping the API key server-side. All AI-generated "
    "content was reviewed by the author; the Reflection in Section 5 is written by the author in their own words."
)
add(
    "European Parliament and Council. (2016) Regulation (EU) 2016/679 on the protection of natural persons with "
    "regard to the processing of personal data (General Data Protection Regulation). OJ L 119, 4.5.2016."
)
add(
    "European Union. (2024) Regulation (EU) 2024/1689 laying down harmonised rules on artificial intelligence "
    "(Artificial Intelligence Act). OJ L, 12.7.2024. Articles 50 (transparency obligations) and Annex III (high-"
    "risk systems)."
)
add(
    "Anthropic. (2026) Claude Sonnet model family: Messages API documentation and tool calling. Available at "
    "https://docs.anthropic.com (Accessed 16 August 2026)."
)
add(
    "CoinGecko. (2026) Public API — /coins/markets. Available at https://www.coingecko.com/en/api (Accessed 16 "
    "August 2026)."
)
add(
    "Alternative.me. (2026) Crypto Fear & Greed Index API. Available at https://alternative.me/crypto/fear-and-"
    "greed-index/ (Accessed 16 August 2026)."
)
add(
    "ExchangeRate-API. (2026) Open endpoint, latest EUR rates (daily-updated). Available at https://open.er-api.com/v6/latest/EUR "
    "(Accessed 16 August 2026)."
)

doc.save(DST)
print("saved:", DST, os.path.getsize(DST), "bytes")