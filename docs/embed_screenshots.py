import os
import re

import docx
from docx.oxml.ns import qn
from docx.shared import Cm

DOC = r"C:\Users\Probook\NCIRL 2025-2026\NCIRL CEAI\CA 3\x25128442_Sydorenko_Nadiya_CEAI_CA3.docx"
SHOTS = r"C:\Users\Probook\NCIRL 2025-2026\NCIRL CEAI\CA 3\screenshots"
WIDTH = Cm(15)

doc = docx.Document(DOC)
placed = 0
for p in list(doc.paragraphs):
    m = re.search(r"INSERT FIGURE (\d+) HERE", p.text)
    if not m:
        continue
    n = int(m.group(1))
    path = os.path.join(SHOTS, f"fig{n}.png")
    if not os.path.exists(path):
        print(f"fig{n}: missing, leaving placeholder")
        continue
    pPr = p._p.get_or_add_pPr()
    for tag in ("w:pBdr", "w:shd"):
        for el in pPr.findall(qn(tag)):
            pPr.remove(el)
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run().add_picture(path, width=WIDTH)
    placed += 1
    print(f"fig{n}: embedded")
doc.save(DOC)
print("placed:", placed, "saved:", DOC)